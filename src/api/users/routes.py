__all__ = ["router"]

import datetime
import io

from docx import Document as create_docx
from docx.document import Document
from docx.shared import Emu
from docx.table import _Cell
from fastapi import APIRouter
from fastapi import Query, Response

from src.api.dependencies import VerifiedDep, VerifiedDepWithUserID
from src.exceptions import ForbiddenException
from src.repositories.users.repository import user_repository
from src.schemas import ViewUser, UserStatus, FillUserProfile
from src.schemas.auth import VerificationSource

router = APIRouter(tags=["Users"])


# docx
@router.get("/users/export", response_class=Response)
async def get_list_of_all_users(verified: VerifiedDep, as_bot: bool = False):
    if as_bot:
        if verified.source != VerificationSource.BOT:
            raise ForbiddenException()
    else:
        if verified.user_id is None:
            raise ForbiddenException()
        issuer = await user_repository.get_user(verified.user_id)
        if issuer.status != UserStatus.LORD:
            raise ForbiddenException()

    users = await user_repository.get_all_users()
    users.sort(key=lambda x: x.name or "")

    document: Document = create_docx()
    document.add_heading("Список музыкальной комнаты", 0)

    table = document.add_table(rows=1, cols=3, style="TableGrid")

    hdr_cells: tuple[_Cell] = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Фамилия Имя"
    hdr_cells[2].text = "Telegram"

    # set bold for header
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for i, user in enumerate(users, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = user.name or "неизвестно"
        row_cells[2].text = f"@{user.alias}"

    # set width for columns
    for cell in table.columns[0].cells:
        cell.width = Emu(500000)

    document.add_page_break()

    bytes_stream = io.BytesIO()
    document.save(bytes_stream)
    bytes_ = bytes_stream.getvalue()

    bytes_stream.close()
    date = datetime.datetime.now().strftime("%d.%m.%Y")
    headers = {
        "Content-Disposition": f"attachment; filename={date}.docx",
    }
    return Response(
        content=bytes_,
        headers=headers,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.get("/users/me")
async def get_me(verified: VerifiedDepWithUserID) -> ViewUser:
    user = await user_repository.get_user(verified.user_id)
    return user


@router.post("/users/me/fill_profile")
async def fill_profile(user: FillUserProfile, verified: VerifiedDepWithUserID) -> ViewUser:
    created = await user_repository.fill_profile(verified.user_id, user)
    return created


@router.get("/users/me/remaining_weekly_hours")
async def get_remaining_weekly_hours(
    verified: VerifiedDep,
    date: datetime.date | None = Query(
        default_factory=datetime.date.today,
        example=datetime.date.today().isoformat(),
        description="Date for which to get remaining hours (iso format). Default: server-side today",
    ),
) -> float:
    start_of_week = date - datetime.timedelta(days=date.weekday())
    ans = await user_repository.remaining_weekly_hours(verified.user_id, start_of_week)
    return ans


@router.get("/users/me/remaining_daily_hours")
async def get_remaining_daily_hours(
    verified: VerifiedDep,
    date: datetime.date | None = Query(
        default_factory=datetime.date.today,
        example=datetime.date.today().isoformat(),
        description="Date for which to get remaining hours (iso format). Default: server-side today",
    ),
) -> float:
    ans = await user_repository.remaining_daily_hours(verified.user_id, date)
    return ans


@router.get("/users/user_id")
async def get_user_id(
    verification: VerifiedDep, telegram_id: int | None = None, email: str | None = None
) -> int | None:
    if verification.source not in (VerificationSource.BOT, VerificationSource.API):
        raise ForbiddenException()
    res = await user_repository.get_user_id(telegram_id=telegram_id, email=email)
    return res
