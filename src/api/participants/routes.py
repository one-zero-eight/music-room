import datetime
import io
from typing import Optional

from docx import Document as create_docx
from docx.document import Document
from docx.shared import Emu
from docx.table import _Cell
from fastapi import Query, Response

from src.api.dependencies import Dependencies, VerifiedDep
from src.api.participants import router
from src.exceptions import ForbiddenException
from src.repositories.participants.abc import AbstractParticipantRepository
from src.schemas import ViewBooking, ViewParticipant, ParticipantStatus, FillParticipantProfile


# docx
@router.get(
    "/export",
    response_class=Response,
)
async def get_list_of_all_users(
    verified: VerifiedDep,
):
    participant_repository = Dependencies.get(AbstractParticipantRepository)
    issuer = await participant_repository.get_participant(verified.user_id)
    if issuer.status != ParticipantStatus.LORD:
        raise ForbiddenException()

    participants = await participant_repository.get_all_participants()

    document: Document = create_docx()
    document.add_heading("Список музыкальной комнаты", 0)

    table = document.add_table(rows=1, cols=3, style="TableGrid")

    hdr_cells: tuple[_Cell] = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Фамилия Имя"
    hdr_cells[2].text = "Telegram"

    # set bold for header
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for i, participant in enumerate(participants, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = participant.name
        row_cells[2].text = f"@{participant.alias}"

    # set width for columns
    for cell in table.columns[0].cells:
        cell.width = Emu(500000)

    document.add_page_break()

    bytes_stream = io.BytesIO()
    document.save(bytes_stream)
    val = bytes_stream.getvalue()
    bytes_stream.close()
    date = datetime.datetime.now().strftime("%d.%m.%Y")
    headers = {
        "Content-Disposition": f"attachment; filename={date}.docx",
    }
    return Response(
        content=val,
        headers=headers,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.put("/{participant_id}/status")
async def change_status(
    participant_id: int,
    new_status: ParticipantStatus,
    verified: VerifiedDep,
) -> ViewParticipant:
    participant_repository = Dependencies.get(AbstractParticipantRepository)

    source = await participant_repository.get_participant(verified.user_id)
    if source.status != ParticipantStatus.LORD:
        raise ForbiddenException()

    updated_participant = await participant_repository.change_status(participant_id, new_status)
    return updated_participant


@router.get("/me")
async def get_me(verified: VerifiedDep) -> ViewParticipant:
    participant_repository = Dependencies.get(AbstractParticipantRepository)
    participant = await participant_repository.get_participant(verified.user_id)
    return participant


@router.post("/me/fill_profile")
async def fill_profile(
    participant: FillParticipantProfile,
    verified: VerifiedDep,
) -> ViewParticipant:
    participant_repository = Dependencies.get(AbstractParticipantRepository)
    created = await participant_repository.fill_profile(verified.user_id, participant)
    return created


@router.get("/me/bookings")
async def get_participant_bookings(verified: VerifiedDep) -> list[ViewBooking]:
    participant_repository = Dependencies.get(AbstractParticipantRepository)
    bookings = await participant_repository.get_participant_bookings(verified.user_id)
    return bookings


@router.get("/me/remaining_weekly_hours")
async def get_remaining_weekly_hours(verified: VerifiedDep) -> float:
    participant_repository = Dependencies.get(AbstractParticipantRepository)
    ans = await participant_repository.remaining_weekly_hours(verified.user_id)
    return ans


@router.get("/me/remaining_daily_hours")
async def get_remaining_daily_hours(
    verified: VerifiedDep,
    date: Optional[datetime.date] = Query(
        default_factory=datetime.date.today,
        example=datetime.date.today().isoformat(),
        description="Date for which to get remaining hours (iso format). Default: server-side today",
    ),
) -> float:
    participant_repository = Dependencies.get(AbstractParticipantRepository)
    ans = await participant_repository.remaining_daily_hours(verified.user_id, date)
    return ans


@router.get("/participant_id")
async def get_participant_id(telegram_id: str) -> int:
    participant_repository = Dependencies.get(AbstractParticipantRepository)

    res = await participant_repository.get_participant_id(telegram_id)
    return res
